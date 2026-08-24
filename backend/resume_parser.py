"""
Turns an uploaded resume file into plain text so it can be sent to the LLM.
Keeping this dumb and dependency-light on purpose: pdfplumber for PDFs,
python-docx for Word files, plain read for .txt.
"""
import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(file_path, ext):
    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext == ".txt":
        return _extract_txt(file_path)
    raise ValueError(f"unsupported extension: {ext}")


def _extract_pdf(file_path):
    chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                chunks.append(text)
    return "\n".join(chunks)


def _extract_docx(file_path):
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # also pull text out of tables (a lot of resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
